#!/usr/bin/env -S uvx --with pdfplumber --with python-docx python
# /// script
# requires-python = ">=3.10"
# dependencies = ["pdfplumber", "python-docx"]
# ///
"""Extract text from CV files (PDF/DOCX)."""

import sys
from pathlib import Path


def extract_pdf(path: Path) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n\n".join(text_parts)


def extract_docx(path: Path) -> str:
    """Extract text from DOCX."""
    from docx import Document
    doc = Document(path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_text(file_path: str) -> str:
    """Extract text from CV file based on extension."""
    path = Path(file_path)
    
    if not path.exists():
        return f"ERROR: File not found: {file_path}"
    
    suffix = path.suffix.lower()
    
    if suffix == ".pdf":
        return extract_pdf(path)
    elif suffix in (".docx", ".doc"):
        if suffix == ".doc":
            return "ERROR: .doc format not supported. Convert to .docx or PDF."
        return extract_docx(path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    else:
        return f"ERROR: Unsupported file format: {suffix}"


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: extract_cv.py <file_path>")
        sys.exit(1)
    
    result = extract_text(sys.argv[1])
    print(result)
